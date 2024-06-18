from fastapi import FastAPI, File, UploadFile, HTTPException
from fastapi.responses import StreamingResponse
import json
import xlsxwriter
import io
from parse_ogrn_nalog import scrape_ogrn_info
from pydantic import BaseModel
from docx import Document


def upload_file(file: UploadFile = File(...)):
    if not file.filename.endswith('.json'):
        raise HTTPException(status_code=400, detail="Unsupported file format")

    contents =file.read()
    users_data = json.loads(contents)

    # Create an Excel file in memory
    output_excel = io.BytesIO()
    workbook = xlsxwriter.Workbook(output_excel, {'in_memory': True})
    worksheet = workbook.add_worksheet()

    # Write headers in the first row
    headers = users_data[0].keys()
    for col, header in enumerate(headers):
        worksheet.write(0, col, header)

    # Write each user's data in the subsequent rows
    for row, user in enumerate(users_data, start=1):
        for col, (key, value) in enumerate(user.items()):
            worksheet.write(row, col, value)

    # Set column widths
    for col, header in enumerate(headers):
        max_width = max(len(str(header)), max(len(str(user[header])) for user in users_data))
        worksheet.set_column(col, col, max_width + 2)  # +2 for padding

    worksheet.autofilter(0, 0, len(users_data), len(headers) - 1)

    workbook.close()
    output_excel.seek(0)

    # Create a Word file in memory
    output_word = io.BytesIO()
    document = Document()

    # Add a heading
    document.add_heading('Users Data', level=1)

    # Create a table with the number of columns equal to the number of headers
    table = document.add_table(rows=1, cols=len(headers))
    hdr_cells = table.rows[0].cells
    for col, header in enumerate(headers):
        hdr_cells[col].text = header

    # Add data to the table
    for user in users_data:
        row_cells = table.add_row().cells
        for col, (key, value) in enumerate(user.items()):
            row_cells[col].text = str(value)

    document.save(output_word)
    output_word.seek(0)

    return {
        "excel_file": StreamingResponse(output_excel, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                        headers={"Content-Disposition": "attachment;filename=all_users_data.xlsx"}),
        "word_file": StreamingResponse(output_word, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       headers={"Content-Disposition": "attachment;filename=all_users_data.docx"})
    }
